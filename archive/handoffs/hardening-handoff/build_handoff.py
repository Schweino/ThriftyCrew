from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Codex\deliverables\hardening-handoff\Thrifty-Crew-Hardening-Integration-Handoff-draft.docx")
SOURCE_URL = "https://docs.google.com/document/d/1rrZh_RwOs5EegfaMQ54tSDD0AHXgZ1HcfveD3rW-6Wg/edit?tab=t.0#heading=h.jz5jxvv3shtw"


def set_font(run, size=None, bold=None, color="000000"):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "DADCE0")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Arial")
    r_fonts.set(qn("w:hAnsi"), "Arial")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_label_paragraph(doc, label, text, style=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(label)
    set_font(r, bold=True)
    r2 = p.add_run(text)
    set_font(r2)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r)
    return p


def add_status_block(doc, status, title, integrated, gap, why):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{status} - {title}")
    set_font(r, size=11, bold=True)
    add_label_paragraph(doc, "Integrated: ", integrated)
    add_label_paragraph(doc, "Not integrated / still open: ", gap)
    add_label_paragraph(doc, "Why: ", why)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0, 0, 0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, before, after, color in (
    ("Heading 1", 20, 20, 6, "000000"),
    ("Heading 2", 16, 18, 6, "000000"),
    ("Heading 3", 14, 16, 4, "434343"),
):
    style = styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(size)
    style.font.bold = False
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Bullet 2", "List Number"):
    style = styles[style_name]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0, 0, 0)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(3)
title.paragraph_format.keep_with_next = True
tr = title.add_run("Thrifty Crew Infrastructure Hardening")
set_font(tr, size=26, bold=False)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(8)
sr = subtitle.add_run("Integration reconciliation and next-step handoff")
set_font(sr, size=14, bold=False, color="434343")

meta = doc.add_paragraph()
meta.paragraph_format.space_after = Pt(12)
mr = meta.add_run("Prepared August 10, 2026 | Production baseline: commit f86121f1")
set_font(mr, size=10, color="555555")

lead = doc.add_paragraph()
lead.paragraph_format.space_after = Pt(12)
lr = lead.add_run("Bottom line: ")
set_font(lr, bold=True)
lr2 = lead.add_run("Most of the plan's control architecture is live. Several operational outcomes are intentionally not claimed because the source plan assumed V3 evidence gates and legacy retirements had already completed; live evidence shows those gates are still accruing. A smaller set of agent-input, per-agent authorization, browser-canary, replication-cutover, archive-query, and source-fix automation items remains genuine implementation work.")
set_font(lr2)

doc.add_heading("1. Purpose and review basis", level=1)
p = doc.add_paragraph("This document compares the linked Infrastructure Hardening Implementation Plan against the repository and production state after phases H0-H6 were implemented. It distinguishes code that exists from proof that has actually been earned. It is written for the next planning agent, not as a marketing completion report.")
p = doc.add_paragraph()
p.add_run("Plan reviewed: ").bold = True
add_hyperlink(p, "Thrifty Crew Infrastructure Hardening - Implementation Plan (V3 Baseline)", SOURCE_URL)
p.add_run(". Repository evidence was reviewed from the hardening status, V3 implementation status, transition inventory, schedule authority, agent registry, workflows, Worker routes, migrations, fixtures, and recovery evidence.")
for run in p.runs:
    set_font(run)

doc.add_heading("2. The baseline mismatch that drove the divergences", level=1)
doc.add_paragraph("The source plan says V3 is fully complete and assumes the PowerShell estate, local watchdog, PC Baker's task, expected-automations.json, and git data bus are retired. That assumption is not supported by the live transition evidence.")
add_bullet(doc, "Shadow ingest: 1 of 14 required clean days.")
add_bullet(doc, "Native semantic parity: 1 of 14 required zero-diff days.")
add_bullet(doc, "Direct Chrome capture: 1 of 4 required weekly cycles.")
add_bullet(doc, "Blind accuracy: 1 of 4 required weekly cycles.")
add_bullet(doc, "Beta releases: 1 of 30 required daily releases and 0 of 4 closed beta weeks.")
add_bullet(doc, "Entitlements: 5 of 7 real lifecycle states; expired and cancelled cannot be fabricated.")
doc.add_paragraph("Because V3 invariant 18 requires evidence-backed cutovers and tested rollback, legacy executors were placed in an explicit transition inventory instead of being tombstoned. This is not unfinished hardening code; it is a refusal to manufacture production history.")

doc.add_heading("3. Phase-level status at a glance", level=1)
table = doc.add_table(rows=1, cols=3)
headers = ["Phase", "Status", "Plain-language result"]
for i, value in enumerate(headers):
    r = table.rows[0].cells[i].paragraphs[0].add_run(value)
    set_font(r, size=10, bold=True)
set_repeat_table_header(table.rows[0])
phase_rows = [
    ("H0", "Partial by evidence", "Truth inventory exists; retirement assumptions are not yet earned."),
    ("H1", "Mostly integrated", "Agents are registered, versioned, budgeted, and ledgered; per-agent enforcement and login canary remain open."),
    ("H2", "Integrated with transition exceptions", "One schedule authority and CI drift gates are live; legacy schedules remain represented until retirement gates close."),
    ("H3", "Platform integrated; operational ports partial", "CI agent runner and one real pilot work; several agents still lack approved live input adapters or chaining."),
    ("H4", "Control path integrated; live wave open", "Immutable content batches and deterministic promotion are live; a full production recipe wave is not yet proven."),
    ("H5", "Mostly integrated with two deliberate redesigns", "Migration and restore controls are complete; replication is canary-only and tiering is capacity-based."),
    ("H6", "Deterministic half integrated", "Source checks and alerts are live; automatic fix PRs, historical replay proof, and final retirements remain open."),
]
for phase, status, result in phase_rows:
    cells = table.add_row().cells
    for i, value in enumerate((phase, status, result)):
        run = cells[i].paragraphs[0].add_run(value)
        set_font(run, size=9.5, bold=(i == 0))
set_table_geometry(table, [900, 2300, 6160])

doc.add_heading("4. Detailed integration reconciliation", level=1)

doc.add_heading("H0 - Verify the baseline", level=2)
add_status_block(doc, "PARTIAL BY LIVE EVIDENCE", "Baseline truth and recovery proof",
                 "A dated transition inventory records every current grocery executor and every open V3 evidence counter. The Worker-only mutation boundary, guard-backed releases, immutable capture path, nightly R2 backup, test corpus, and production restore drill are implemented and exercised.",
                 "The plan's retirement check is not green. Multiple legacy PC and GitHub jobs still execute under transition status, and the V3 soak counters remain open.",
                 "The plan's baseline was asserted more strongly than production evidence allows. Retiring rollback paths before 14-day, 4-week, 30-day, and entitlement gates close would violate the V3 cutover rule.")

doc.add_heading("H1 - Agents become code", level=2)
add_status_block(doc, "MOSTLY INTEGRATED", "A-1 agent prompts and A-2 registry",
                 "Ten agents are declared in config/agents.json. Prompts live under platform/agents, hashes are checked, models and fallbacks are pinned, token prices are effective-dated, budgets and fixture files are declared, and registry drift is part of CI.",
                 "The implementation proves the declared registry; it does not prove that every former desktop-app registration has been retired.",
                 "Desktop retirement is coupled to real scheduled successor runs, not merely to a registry row.")
add_status_block(doc, "INTEGRATED", "A-3 ledger and budget accounting",
                 "Agent job starts validate prompt/model identity, reserve budget, and finish with token and micro-dollar usage. Optional agents fail closed when their budget is exhausted. If the ledger cannot start, the workflow permits only diagnostic output and then fails.",
                 "The status page does not yet present a mature operator-facing budget dashboard for every agent, although the durable data and authorization API exist.",
                 "The control was prioritized before presentation. The missing presentation is product/operations polish rather than a write-boundary gap.")
add_status_block(doc, "PARTIAL", "A-5 scoped credentials",
                 "GitHub OIDC is bound to the exact repository and three exact workflow files. The agent workflow is route-family restricted; the restore workflow can only trigger a restore drill. HMAC roles still protect deterministic clients.",
                 "The source plan asked for cryptographically distinct per-agent credentials. Current CI authentication identifies the GitHub workflow run, while per-agent capabilities are registry declarations; the Worker does not yet enforce each capability against every endpoint for the named agent.",
                 "Workflow-scoped OIDC removed static CI secrets and sharply reduced blast radius, but it is not equivalent to per-agent authorization. This remains real security hardening work.")
add_status_block(doc, "NOT INTEGRATED", "Store-login canary",
                 "No Tuesday store-session canary schedule or store-specific signed-out acceptance proof exists.",
                 "The task is absent from config/schedules.json and from the PC executor inventory.",
                 "A reliable canary needs store-specific authenticated Chrome probes and a deliberate sign-out test. That could not be safely simulated during infrastructure work and should be implemented as a deterministic browser client, not an AI agent.")

doc.add_heading("H2 - One schedule source", level=2)
add_status_block(doc, "INTEGRATED", "S-1 schedule authority and drift gate",
                 "platform/config/schedules.json owns active and transition schedules with timezone, executor, owner, proof, maximum gap, lifecycle, and retirement gate. tc schedules check performs bidirectional verification against GitHub workflow crons, Worker cron, and grocery/expected-automations.json. Rogue or missing entries fail CI.",
                 "The plan expected expected-automations.json and several PC definitions to disappear immediately. They remain as verified transition surfaces.",
                 "Keeping a transition executor visible to the same authority prevents split-brain schedules while its successor earns evidence. Deleting the row would hide rather than retire the job.")
add_status_block(doc, "PARTIAL", "S-2 shared job design standard",
                 "Cloud jobs use bounded timeouts, durable start/finish ledger rows, honest failures, idempotency, concurrency controls, and recovery dispatch. CLAUDE.md and the implementation contract encode the safety rules.",
                 "There is no single reusable job-harness package and no separately named pull-request checklist artifact covering every legacy and future executor.",
                 "The common behavior is enforced in the current workflows and APIs. Extracting it into a universal harness is useful cleanup, but was not required to close the observed failure modes.")

doc.add_heading("H3 - Plane split", level=2)
add_status_block(doc, "PARTIAL", "Judgment agents in CI",
                 "The @openai/agents runner executes registered CI agents, validates prompt hashes and approved models, records usage, and retains bounded artifacts. The post-publish reviewer completed a real scheduled-style GitHub run with a durable ledger. No judgment agent is assigned to the PC plane.",
                 "The triage reviewer and recipe sourcer currently produce a typed no-op when no approved input file is supplied. The accuracy-headless and triage-developer agents are registered but not independently scheduled. The reviewer-to-developer PR chain is not automatic.",
                 "The implementation deliberately fails closed instead of letting a scheduled agent invent input or mutate from freeform text. Real adapters from triage_items, accuracy draws, and content requests must be built before those ports can be claimed.")
add_status_block(doc, "NOT YET AT TARGET STATE", "PC workload reduction",
                 "The judgment-agent registry contains no PC agents, which meets the architectural direction.",
                 "The overall PC estate is not yet the source plan's exact two workloads plus two tasks; the transition inventory still lists legacy Baker's, Family Fare, daily pipeline, watchdog, wake, and Friday-email jobs.",
                 "Those are rollback and evidence-accrual executors. Their removal is gated by production history, not by the hardening implementation calendar.")

doc.add_heading("H4 - Content batches", level=2)
add_status_block(doc, "CONTROL PATH INTEGRATED; LIVE WAVE OPEN", "Immutable recipe staging and deterministic promotion",
                 "D1 content batches accept immutable staged items, record an agent audit, and require a separate deterministic promotion call. Guards check duplicate identity, commodity mapping, ingredient coherence, provenance, and content hashes. Rejected or promoted items cannot be edited. Passing and failing fixtures exercise the gate.",
                 "The scheduled recipe-pack workflow does not yet orchestrate sourcer, deduper, mapper, writer, and auditor through real approved inputs. No production evidence shows one complete recipe wave plus deliberate bad-spec and failed-batch trials.",
                 "The safety boundary was built first. Claiming the full recipe operation without structured inter-agent adapters and a live wave would confuse API readiness with production completion.")

doc.add_heading("H5 - Database finishing moves", level=2)
add_status_block(doc, "INTENTIONALLY CANARY-ONLY", "D-1 D1 Sessions read replication",
                 "An isolated /api/v2/replica-canary route uses a first-primary Sessions API strategy. It allows consistency and benefit measurements without changing public release semantics.",
                 "Public /v2 reads have not been moved to replicas and a full week of zero-consistency-finding traffic has not accrued.",
                 "The source plan made replication a build step and its safety gate a later check. The implementation reverses that order: prove the canary before exposing public reads. This is deliberate risk control, not an accidental omission.")
add_status_block(doc, "INTEGRATED WITH A POLICY CHANGE", "D-2 migration governance",
                 "Migration 0016 is additive. CI enforces expand-contract, forbids foreign_keys=off, rejects nonconforming migrations, and requires durable restore proof before a future contract migration.",
                 "The plan requested a down path or documented no-down rationale. The implemented policy forbids down-migration files and uses forward fixes plus restore evidence.",
                 "D1 production rollbacks through reverse DDL can be lossy and misleading. Forward-only expand/contract with tested restoration is the safer operational policy.")
add_status_block(doc, "FULLY INTEGRATED AND PROVEN", "D-4 automated restore drills",
                 "A quarterly GitHub workflow uses exact-workflow OIDC to trigger D1RestoreDrillWorkflow. It selects the promoted backup batch first, restores into a named scratch D1, verifies six backup-derived table counts plus release ID/hash, records durable evidence, and removes exact scratch and staging objects.",
                 "No implementation gap remains. Future quarters still need scheduled successes, as required by the invariant.",
                 "Production drill d1-restore-2026-Q3-a20 passed on August 10, 2026, recovered ten oversized release payloads, and left zero scratch databases.")
add_status_block(doc, "INTEGRATED WITH A BETTER TRIGGER; QUERY TIER OPEN", "D-3 cold tiering",
                 "Daily forecasts use actual D1 bytes, measured growth, the configured 10 GiB limit, protected release references, and projected exhaustion. Execution is disarmed below 70% utilization. When armed, only unreferenced observations older than 18 months can enter a verified, SHA-256-addressed Parquet manifest in R2.",
                 "The fixed 5 million-row trigger was not implemented. No live deletion/archive cycle has fired while capacity is healthy, and a DuckDB analytical query service has not been activated.",
                 "Row count is a poor proxy for D1 storage because row width and indexes matter. Building an archive query plane before an actual analytical workload and before the threshold fires would add unused infrastructure. The trigger was replaced, not forgotten; the query consumer remains next-step work when needed.")

doc.add_heading("H6 - Sentinel and closeout", level=2)
add_status_block(doc, "PARTIAL", "A-8 source-contract sentinel",
                 "Deterministic source contracts run before server-source ingestion and cover row presence, term completion, taxonomy availability, and price mode. Failures write durable sentinel results and operational alerts. A source-sentinel investigator prompt and registry entry exist as a read-only second stage.",
                 "The scheduled path currently runs deterministic tc sentinel latest rather than an AI investigation that opens a pull request. No historical source-change replay has produced a verified PR draft.",
                 "A source-shape check must be deterministic and block ingestion immediately. Granting an AI agent repository write access before its inputs, diff boundaries, tests, and PR-only token are proven would weaken the system. The investigative PR stage remains a bounded follow-up.")
add_status_block(doc, "NOT YET AUTHORIZED", "Final retirements",
                 "Every legacy executor has a named successor and retirement gate in the transition inventory.",
                 "Desktop scheduler entries and legacy publication/capture tasks have not been tombstoned.",
                 "The V3 evidence counters are still open. Retiring them now would remove the tested rollback path before the approved soak is complete.")

doc.add_heading("5. Intentional pushbacks and replacements", level=1)
add_number(doc, "Replication was staged behind a canary instead of switching public reads immediately. The public release pointer remains primary-authoritative until a week of evidence proves safety and value.")
add_number(doc, "Down migrations were rejected in favor of forward-only expand/contract and durable restore proof. This avoids false confidence in lossy reverse DDL.")
add_number(doc, "The 5 million-row archive trigger was replaced with byte utilization, measured growth, protected references, and projected exhaustion. Capacity should be measured in the unit the platform limits.")
add_number(doc, "AI audit was not allowed to promote recipe content. The agent may report findings; deterministic Worker guards own promotion.")
add_number(doc, "Legacy executors were not retired on a calendar assertion. They remain visible, verified, and explicitly transitional until their production evidence gates close.")
add_number(doc, "A missing agent ledger or approved input does not fall back to invisible work. It produces bounded diagnostics or a typed no-op and does not authorize mutation.")

doc.add_heading("6. Definition-of-done crosswalk", level=1)
dod = doc.add_table(rows=1, cols=3)
for i, value in enumerate(("Source-plan outcome", "Status", "What remains")):
    r = dod.rows[0].cells[i].paragraphs[0].add_run(value)
    set_font(r, size=10, bold=True)
set_repeat_table_header(dod.rows[0])
dod_rows = [
    ("1. One source for every recurring job and agent", "Partial", "Authority and drift gates are live, but transition executors are still active by design."),
    ("2. Prompts, registry, ledger, token cost, least privilege", "Partial", "Prompts/registry/ledger/cost are live; cryptographic per-agent endpoint scopes remain open."),
    ("3. PC reduced to the exact target footprint", "Not met", "Wait for V3 retirement gates, add login canary, then tombstone named legacy tasks."),
    ("4. Recipe pack only through staged promoted batches", "Control complete; operation open", "Run and prove a full five-agent production wave with negative tests."),
    ("5. Site change yields same-day sentinel PR and triage", "Partial", "Deterministic failure and triage exist; scoped PR generation and historical replay proof remain."),
    ("6. Replication, migrations, restore, tiering, archive analytics", "Partial", "Migration/restore/tiering controls are live; public replication and archive query consumer await gates/use."),
    ("7. All H0 findings closed", "Not met", "Time-based V3 evidence counters and real entitlement transitions remain open."),
]
for outcome, status, remains in dod_rows:
    cells = dod.add_row().cells
    for i, value in enumerate((outcome, status, remains)):
        run = cells[i].paragraphs[0].add_run(value)
        set_font(run, size=9.25, bold=(i == 1))
set_table_geometry(dod, [3800, 1500, 4060])

doc.add_heading("7. Recommended next-step backlog", level=1)
doc.add_heading("Priority 0 - Close genuine implementation gaps", level=2)
add_number(doc, "Build approved-input adapters for triage_items, accuracy draws, recipe requests, and content-batch state. Validate each input contract before invoking a model.")
add_number(doc, "Enforce registry capabilities per named agent at the Worker, not only per workflow route family. Use short-lived OIDC-derived identity or an equivalent signed claim; reject a deliberately over-scoped call for every agent class and persist the audit event.")
add_number(doc, "Complete reviewer-to-developer chaining with typed artifacts and a PR-only repository credential. The developer must never receive production mutation scope.")
add_number(doc, "Implement the Tuesday deterministic Chrome store-login canary, cloud receipt, and deliberate signed-out acceptance test.")
add_number(doc, "Run one real five-stage recipe wave through immutable content batches. Include one schema-invalid item and one audited-but-deterministically-rejected batch.")
add_number(doc, "Add the source-sentinel investigator's bounded evidence bundle, extraction-diff generator, tests, and PR-only submission. Replay at least one archived source change end to end.")

doc.add_heading("Priority 1 - Earn and execute the gated cutovers", level=2)
add_number(doc, "Accrue replica-canary latency and consistency evidence for seven days. Move one low-risk public read only if the benefit is measurable and consistency findings remain zero; retain an immediate primary rollback.")
add_number(doc, "Continue V3 evidence accrual: 14 shadow days, 14 parity days, four Chrome weeks, four accuracy weeks, 30 beta releases, four beta weeks, and the remaining real entitlement states.")
add_number(doc, "Retire each legacy executor only when its own transition-inventory gate closes. Commit a tombstone naming the successor, rollback window, and evidence reference.")
add_number(doc, "When analytical demand exists, add a read-only DuckDB/R2 query path and prove it against a fixture archive. Do not enable archive deletion while D1 is below the 70% capacity gate.")

doc.add_heading("Priority 2 - Consolidation and operator experience", level=2)
add_number(doc, "Extract the repeated ledger/timeout/idempotency behavior into a shared job harness and add a concise pull-request checklist.")
add_number(doc, "Expose agent spend versus budget, skipped/no-op reason, prompt hash, and last successful live-input run on the operator status surface.")
add_number(doc, "Upgrade deprecated GitHub action runtime dependencies noted by Actions so warning-free CI remains part of the baseline.")

doc.add_heading("8. Acceptance gates for the next agent", level=1)
add_bullet(doc, "Do not mark an agent port complete when it only ran a fixture or a no-op. Require a real approved input, a durable ledger row, token cost, typed output, and a rejected over-scope test.")
add_bullet(doc, "Do not mark recipe operations complete when only the staging API is green. Require one real full wave and negative promotion tests.")
add_bullet(doc, "Do not move public reads to replication without seven days of canary evidence and an exercised primary rollback.")
add_bullet(doc, "Do not delete legacy schedules merely to make the inventory look clean. Close the named evidence gate first, then tombstone the exact executor.")
add_bullet(doc, "Do not grant the sentinel or triage developer a production mutation credential. PR-only repository scope is the ceiling.")
add_bullet(doc, "Do not fabricate expired/cancelled member states, weekly cycles, daily releases, or source-change incidents. Calendar and lifecycle evidence must be real.")

doc.add_heading("9. Production evidence and repository references", level=1)
add_bullet(doc, "Final documentation commit: f86121f1. GitHub platform deployment and gates completed successfully.")
add_bullet(doc, "Production restore drill: d1-restore-2026-Q3-a20, passed August 10, 2026; six table counts and release ID/hash matched; ten oversized payloads recovered; zero scratch databases remained.")
add_bullet(doc, "Automated checks: 102 passed, 1 intentional skip; typechecks, builds, migrations, configuration, schedules, agents, and Worker dry run passed.")
add_bullet(doc, "Real agent pilot: GitHub Actions run 31361101543 completed the post-publish reviewer workflow with durable evidence.")
add_bullet(doc, "Core evidence files: platform/docs/HARDENING-STATUS.md, platform/docs/IMPLEMENTATION-STATUS.md, platform/docs/IMPLEMENTATION-CONTRACT.md, platform/config/transition-inventory.json, platform/config/schedules.json, platform/config/agents.json, .github/workflows/platform-agents.yml, .github/workflows/platform-restore.yml.")

doc.add_heading("10. Planning conclusion", level=1)
doc.add_paragraph("The hardening work should be treated as a strong control-plane foundation, not as proof that every agent and retirement gate in the source plan is operationally complete. The next plan should focus on real-input agent adapters, enforceable per-agent authorization, browser-session canaries, full content-wave proof, sentinel PR automation, and evidence-gated cutovers. It should not rebuild the schedule authority, content-batch safety boundary, migration policy, restore workflow, capacity forecast, or deterministic source contracts that are already deployed.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
